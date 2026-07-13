from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .model import ReportDocument
from .theme import COBALT_HEX, MUTED_HEX, hex_rgb


_COBALT = RGBColor(*hex_rgb(COBALT_HEX))
_MUTED = RGBColor(*hex_rgb(MUTED_HEX))


def generate_docx(report: ReportDocument) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size in (("Title", 22), ("Heading 1", 16), ("Heading 2", 13)):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = _COBALT
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    header = section.header.paragraphs[0]
    header.text = "LUMINA · AGENT REPORT"
    header.runs[0].font.color.rgb = _COBALT
    header.runs[0].font.size = Pt(9)

    title = document.add_paragraph(style="Title")
    title.add_run(report.title)
    title.paragraph_format.space_after = Pt(8)
    summary = document.add_paragraph(report.executive_summary)
    summary.paragraph_format.space_after = Pt(16)

    document.add_heading("핵심 지표", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Light Shading Accent 1"
    table.rows[0].cells[0].text = "지표"
    table.rows[0].cells[1].text = "값"
    for metric in report.metrics:
        cells = table.add_row().cells
        cells[0].text = metric.label
        cells[1].text = metric.value

    document.add_heading("원 요청", level=1)
    request = document.add_paragraph(report.request)
    request.paragraph_format.left_indent = Inches(0.2)
    request.paragraph_format.right_indent = Inches(0.2)
    request.paragraph_format.space_after = Pt(12)

    for report_section in report.sections:
        document.add_heading(report_section.heading, level=1)
        if report_section.body:
            document.add_paragraph(report_section.body)
        for bullet in report_section.bullets:
            document.add_paragraph(bullet, style="List Bullet")

    if report.action_items:
        document.add_heading("후속 조치", level=1)
        for action in report.action_items:
            document.add_paragraph(action, style="List Number")

    if report.images:
        document.add_heading("이미지 자산", level=1)
        for image in report.images:
            try:
                document.add_picture(BytesIO(image.content), width=Inches(6.2))
            except (OSError, ValueError) as exc:
                raise ValueError(
                    f"DOCX에 삽입할 수 없는 이미지 형식입니다: {image.display_name}"
                ) from exc
            caption = document.add_paragraph(image.display_name)
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Lumina가 생성한 편집 가능한 Artifact")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = _MUTED

    output = BytesIO()
    document.save(output)
    return output.getvalue()
